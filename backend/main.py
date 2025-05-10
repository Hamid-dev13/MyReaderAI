from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import tempfile
import os
import requests
from extract_placeholders import extract_placeholders_from_docx
from extract_placeholders_advanced import extract_placeholders_with_context, analyze_placeholder_types
# Remplacer PyPDF2 par pdfminer.six
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from typing import Dict, Any, Optional
from pydantic import BaseModel
from fastapi.responses import StreamingResponse
import io

app = FastAPI(title="RAEDIFICARE Template API")

# URL du webhook n8n
N8N_WEBHOOK_URL = "https://hamiddev13.app.n8n.cloud/webhook-test/raedificare-template"

# Ajouter CORS pour permettre les requêtes depuis le frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Pour le développement, à restreindre en production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Fonction d'extraction de texte PDF améliorée avec pdfminer.six
def extract_text_from_pdf(file_path):
    """
    Extrait tout le texte d'un fichier PDF en utilisant pdfminer.six
    """
    try:
        if not os.path.exists(file_path):
            return {"error": "Le fichier n'existe pas"}
            
        # Utiliser pdfminer avec des paramètres optimisés pour une meilleure extraction
        laparams = LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            all_texts=True
        )
        
        # Extraire tout le texte du PDF
        full_text = extract_text(
            file_path,
            laparams=laparams,
            codec='utf-8'
        )
        
        # Compter approximativement le nombre de pages (basé sur les sauts de page)
        page_breaks = full_text.count('\f') + 1
        
        return {
            "success": True,
            "total_pages": page_breaks,
            "text": full_text
        }
            
    except Exception as e:
        return {
            "success": False,
            "error": f"Erreur lors de l'extraction du texte: {str(e)}"
        }

@app.get("/")
def read_root():
    return {"message": "Bienvenue sur l'API RAEDIFICARE Template"}

@app.post("/extract-placeholders/")
async def extract_placeholders(file: UploadFile = File(...)):
    """
    Extrait tous les placeholders d'un fichier template docx uploadé
    """
    # Vérifier que c'est bien un fichier docx
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format .docx")
    
    # Sauvegarder temporairement le fichier uploadé
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_file_path = temp_file.name
    
    try:
        # Écrire le contenu du fichier uploadé dans le fichier temporaire
        content = await file.read()
        with open(temp_file_path, "wb") as f:
            f.write(content)
        
        # Extraire les placeholders
        result = extract_placeholders_from_docx(temp_file_path)
        
        # Vérifier s'il y a eu une erreur
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["error"])
            
        return result
        
    finally:
        # Nettoyer le fichier temporaire
        temp_file.close()
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

@app.post("/analyze-template-advanced/")
async def analyze_template_advanced(file: UploadFile = File(...)):
    """
    Analyse avancée d'un template avec contexte des placeholders
    """
    # Vérifier que c'est bien un fichier docx
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format .docx")
    
    # Sauvegarder temporairement le fichier uploadé
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_file_path = temp_file.name
    
    try:
        # Écrire le contenu du fichier uploadé dans le fichier temporaire
        content = await file.read()
        with open(temp_file_path, "wb") as f:
            f.write(content)
        
        # Extraire les placeholders avec contexte
        placeholders = extract_placeholders_with_context(temp_file_path)
        
        # Vérifier s'il y a eu une erreur
        if "error" in placeholders:
            raise HTTPException(status_code=500, detail=placeholders["error"])
        
        # Analyser les types de placeholders
        placeholders_with_types = analyze_placeholder_types(placeholders)
        
        return placeholders_with_types
        
    finally:
        # Nettoyer le fichier temporaire
        temp_file.close()
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

@app.post("/send-to-n8n/")
async def send_to_n8n(file: UploadFile = File(...)):
    """
    Extrait les placeholders d'un template et les envoie à n8n
    """
    # Vérifier que c'est bien un fichier docx
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format .docx")
    
    # Sauvegarder temporairement le fichier uploadé
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_file_path = temp_file.name
    
    try:
        # Écrire le contenu du fichier uploadé dans le fichier temporaire
        content = await file.read()
        with open(temp_file_path, "wb") as f:
            f.write(content)
        
        # Extraire les placeholders avec contexte
        result = extract_placeholders_with_context(temp_file_path)
        
        # Vérifier s'il y a eu une erreur
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["error"])
        
        # Analyser les types
        result = analyze_placeholder_types(result)
        
        # Préparer les données pour n8n
        placeholders_list = [ph["placeholder"] for ph in result["unique_placeholders"]]
        data_for_n8n = {
            "filename": file.filename,
            "total_found": result["total_found"],
            "unique_count": result["unique_count"],
            "placeholders": placeholders_list
        }
        
        # Envoyer à n8n
        try:
            n8n_response = requests.post(
                N8N_WEBHOOK_URL,
                json=data_for_n8n,
                headers={"Content-Type": "application/json"}
            )
            
            if n8n_response.status_code != 200:
                return {
                    "n8n_status": "error",
                    "status_code": n8n_response.status_code,
                    "error": f"Erreur n8n: {n8n_response.text}",
                    "placeholders_data": data_for_n8n
                }
                
            return {
                "n8n_status": "success",
                "status_code": n8n_response.status_code,
                "n8n_response": n8n_response.text,
                "placeholders_data": data_for_n8n
            }
            
        except Exception as e:
            return {
                "n8n_status": "error",
                "error": f"Erreur de connexion à n8n: {str(e)}",
                "placeholders_data": data_for_n8n
            }
        
    finally:
        # Nettoyer le fichier temporaire
        temp_file.close()
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

# NOUVEL ENDPOINT POUR EXTRAIRE LE TEXTE D'UN PDF
@app.post("/extract-pdf-text/")
async def extract_pdf_text(file: UploadFile = File(...)):
    """
    Extrait tout le texte d'un fichier PDF uploadé
    """
    # Vérifier que c'est bien un fichier PDF
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Le fichier doit être au format PDF")
    
    # Sauvegarder temporairement le fichier uploadé
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_file_path = temp_file.name
    
    try:
        # Écrire le contenu du fichier uploadé dans le fichier temporaire
        content = await file.read()
        with open(temp_file_path, "wb") as f:
            f.write(content)
        
        # Extraire le texte du PDF avec pdfminer.six
        result = extract_text_from_pdf(temp_file_path)
        
        # Vérifier s'il y a eu une erreur
        if "error" in result:
            raise HTTPException(status_code=500, detail=result["error"])
            
        return result
        
    finally:
        # Nettoyer le fichier temporaire
        temp_file.close()
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

# ENDPOINT POUR TESTER LA CONNEXION À N8N
@app.post("/test-n8n-connection/")
async def test_n8n_connection():
    """
    Teste la connexion au webhook n8n avec des données de test
    """
    test_data = {
        "filename": "test_swagger.docx",
        "placeholders": ["test_placeholder_1", "test_placeholder_2", "test_placeholder_3"],
        "message": "Test depuis l'API Swagger"
    }
    
    try:
        n8n_response = requests.post(
            N8N_WEBHOOK_URL,
            json=test_data,
            headers={"Content-Type": "application/json"}
        )
        
        return {
            "n8n_status": "success" if n8n_response.status_code == 200 else "error",
            "status_code": n8n_response.status_code,
            "response": n8n_response.text
        }
        
    except Exception as e:
        return {
            "n8n_status": "error",
            "error": str(e)
        }

# MODÈLE PYDANTIC POUR LA REQUÊTE DE TRAITEMENT V3
class ProcessTextForV3Request(BaseModel):
    text: str
    current_data: Optional[Dict[str, str]] = {}

# NOUVEL ENDPOINT POUR TRAITER LE TEXTE ET EXTRAIRE LES DONNÉES V3
@app.post("/process-text-for-v3/")
async def process_text_for_v3(request: ProcessTextForV3Request):
    """
    Traite le texte extrait d'un PDF pour en extraire des données structurées pour le document V3
    """
    try:
        text = request.text
        current_data = request.current_data or {}
        
        if not text:
            return {
                "success": False,
                "error": "Aucun texte fourni pour le traitement"
            }
        
        # Dictionnaire pour stocker les données extraites
        extracted_data = {}
        
        # Logique d'extraction simple pour les champs courants
        # Ces règles d'extraction peuvent être améliorées selon les besoins
        
        # Analyse des coordonnées du chantier
        if "Adresse" in text:
            # Recherche de l'adresse après le mot "Adresse"
            import re
            address_match = re.search(r'Adresse[:\s]*([^\n]+)', text)
            if address_match:
                extracted_data["adresse_chantier"] = address_match.group(1).strip()
        
        # Analyse du code postal
        postal_code_match = re.search(r'\b(\d{5})\b', text)
        if postal_code_match:
            extracted_data["code_postal"] = postal_code_match.group(1)
        
        # Analyse de la ville
        if "Ville" in text:
            city_match = re.search(r'Ville[:\s]*([^\n]+)', text)
            if city_match:
                extracted_data["ville"] = city_match.group(1).strip()
        
        # Analyse du maître d'ouvrage
        if "Maître d'ouvrage" in text or "Maitre d'ouvrage" in text:
            mo_match = re.search(r'Ma[îi]tre d\'ouvrage[:\s]*([^\n]+)', text)
            if mo_match:
                extracted_data["maitre_ouvrage"] = mo_match.group(1).strip()
        
        # Analyse de la référence du projet
        if "Référence" in text:
            ref_match = re.search(r'Référence[:\s]*([^\n]+)', text)
            if ref_match:
                extracted_data["reference_projet"] = ref_match.group(1).strip()
        
        # Analyse du type de bâtiment
        batiment_types = ["Immeuble", "Maison", "Bureau", "Entrepôt", "Usine", "Commercial"]
        for bt in batiment_types:
            if bt in text:
                extracted_data["type_batiment"] = bt
                break
        
        # Fusionner avec les données actuelles (les nouvelles données ont priorité)
        merged_data = {**current_data, **extracted_data}
        
        return {
            "success": True,
            "data": merged_data
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Erreur lors du traitement du texte pour V3: {str(e)}"
        }

class DocxGenerationRequest(BaseModel):
    values: Dict[str, str]

@app.post("/generate-docx/")
async def generate_docx(file: UploadFile = File(...), json_values: str = File(...)):
    import json
    from docx import Document
    import tempfile
    import re
    from pathlib import Path
    import io
    import os

    # Charger les valeurs du JSON
    try:
        data = json.loads(json_values)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"JSON invalide: {str(e)}")

    # Logique de choix selon le booléen rapport_type
    rapport_type = data.get("rapport_type", False)
    if rapport_type is True:
        values = {
            "rapport_amiante_type_1": "Ces rapports ne font pas état de présence d'amiante dans les matériaux présents sur les emprises du présent diagnostic PEMD Ressources.",
            "rapport_amiante_type_2": ""
        }
    else:
        values = {
            "rapport_amiante_type_1": "",
            "rapport_amiante_type_2": "Ces rapports font état de présence d'amiante, mais hors des emprises concernées par le présent diagnostic PEMD Ressources : préciser les matériaux amiantés. Voir détails dans les conclusions des diagnostics en question"
        }

    # Logique de choix selon le booléen rapport_plomb_type
    rapport_plomb_type = data.get("rapport_plomb_type", False)
    if rapport_plomb_type is True:
        values["rapport_plomb_type_true"] = "Ces rapports ne font pas état de présence de plomb dans les matériaux présents sur les emprises du présent diagnostic Ressources"
        values["rapport_plomb_type_false"] = ""
    else:
        values["rapport_plomb_type_true"] = ""
        values["rapport_plomb_type_false"] = "Ces rapports font état de présence de plomb, mais hors des emprises concernées par le présent diagnostic PEM-Ressources : préciser les matériaux plombés. Voir détails dans les conclusions des diagnostics en question."

    # Logique de choix selon le booléen rapport_type_termites
    rapport_type_termites = data.get("rapport_type_termites", False)
    if rapport_type_termites is True:
        values["rapport_type_termites_true"] = "Ces rapports ne font pas état de la présence de termites dans les matériaux présents sur les emprises projet."
        values["rapport_type_termites_false"] = ""
    else:
        values["rapport_type_termites_true"] = ""
        values["rapport_type_termites_false"] = "Ces rapports font état de présence de termites, mais hors des emprises concernées par le présent diagnostic Ressources : préciser les matériaux concernés. Voir détails dans les conclusions des diagnostics en question."

    # Logique dynamique pour rapport enrobés/étanchéités
    pollution_type = data.get("pollution_type", "aucun")
    support_type = data.get("support_type", "enrobes")

    # Construction du texte pour le support
    if support_type == "enrobes":
        support_txt = "enrobés"
    elif support_type == "etancheites":
        support_txt = "étanchéités"
    elif support_type == "les_deux":
        support_txt = "enrobés / étanchéités"
    else:
        support_txt = "enrobés"

    # Construction du texte pour la pollution
    if pollution_type == "aucun":
        values["rapport_type_enrobes"] = (
            f"Ces rapports ne font pas état de présence d'amiante ou HAP dans les {support_txt} présentes sur les emprises projet."
        )
    elif pollution_type == "amiante":
        values["rapport_type_enrobes"] = (
            f"Ces rapports font état de présence d'amiante dans les {support_txt} sur les emprises projet. Voir détails dans les conclusions des diagnostics en question."
        )
    elif pollution_type == "hap":
        values["rapport_type_enrobes"] = (
            f"Ces rapports font état de présence de pollution HAP dans les {support_txt} sur les emprises projet. Voir détails dans les conclusions des diagnostics en question."
        )
    elif pollution_type == "les_deux":
        values["rapport_type_enrobes"] = (
            f"Ces rapports font état de présence d'amiante et de pollution HAP dans les {support_txt} sur les emprises projet. Voir détails dans les conclusions des diagnostics en question."
        )
    else:
        values["rapport_type_enrobes"] = ""

    # Logique de choix pour diagnostic de pollution
    diagnostic_de_pollution = data.get("diagnostic_de_pollution", False)
    if diagnostic_de_pollution is True:
        values["diagnostic_de_pollution_true"] = "Ces derniers ont été diagnostiqués dans le diagnostic de pollution."
        values["diagnostic_de_pollution_false"] = ""
    else:
        values["diagnostic_de_pollution_true"] = ""
        values["diagnostic_de_pollution_false"] = "Ces derniers n'ont pas été diagnostiqués dans le diagnostic de pollution."

    # Logique de choix pour site occupé
    site_occupé = data.get("site_occupé", False)
    if site_occupé is True:
        values["site_occupé_true"] = "occupé"
        values["site_occupé_false"] = ""
    else:
        values["site_occupé_true"] = ""
        values["site_occupé_false"] = "non occupé"

    # Ajouter les autres champs éventuels du JSON
    for k, v in data.items():
        if k not in values:
            values[k] = v

    # Création d'un pattern regex unique pour tous les placeholders
    placeholder_pattern = re.compile(r'\$\{(' + '|'.join(re.escape(k) for k in values.keys()) + r')\}')
    
    # Utilisation d'un context manager pour le fichier temporaire
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_template:
        temp_template_path = temp_template.name
        try:
            # Écriture du contenu
            content = await file.read()
            with open(temp_template_path, "wb") as f:
                f.write(content)
            
            # Création du document
            doc = Document(temp_template_path)
            
            # Remplacement dans les paragraphes
            for paragraph in doc.paragraphs:
                if placeholder_pattern.search(paragraph.text):
                    paragraph.text = placeholder_pattern.sub(
                        lambda m: str(values.get(m.group(1), "")),
                        paragraph.text
                    )
            
            # Remplacement dans les tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder_pattern.search(paragraph.text):
                                paragraph.text = placeholder_pattern.sub(
                                    lambda m: str(values.get(m.group(1), "")),
                                    paragraph.text
                                )
            
            # Sauvegarde dans un buffer mémoire
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            
            # Nettoyage du nom de fichier
            safe_filename = Path(file.filename).stem
            safe_filename = re.sub(r'[^\w\-_.]', '_', safe_filename)
            
            return StreamingResponse(
                output_stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=generated_{safe_filename}.docx"
                }
            )
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Erreur lors de la génération du DOCX: {str(e)}"
            )
        finally:
            # Nettoyage du fichier temporaire
            if os.path.exists(temp_template_path):
                os.unlink(temp_template_path)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)