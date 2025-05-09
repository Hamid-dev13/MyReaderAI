import re
from docx import Document
import os

def extract_placeholders_from_docx(docx_path):
    """
    Extrait tous les placeholders de type ${} d'un document Word
    
    Args:
        docx_path (str): Chemin vers le fichier docx
        
    Returns:
        list: Liste des placeholders uniques trouvés
    """
    try:
        # Vérifier si le fichier existe
        if not os.path.exists(docx_path):
            return {"error": f"Le fichier {docx_path} n'existe pas"}
            
        # Ouvrir le document Word
        doc = Document(docx_path)
        
        # Extraire tout le texte
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
            
        # Extraire également le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text += paragraph.text + "\n"
        
        # Utiliser une expression régulière pour trouver tous les placeholders
        pattern = r'\${([^}]*)}'
        placeholders = re.findall(pattern, full_text)
        
        # Éliminer les doublons
        unique_placeholders = list(set(placeholders))
        
        return {
            "total_found": len(placeholders),
            "unique_placeholders": unique_placeholders
        }
        
    except Exception as e:
        return {"error": str(e)}

# Test de la fonction
if __name__ == "__main__":
    # Remplacer par le chemin de votre template
    template_path = "chemin/vers/votre/template.docx"
    result = extract_placeholders_from_docx(template_path)
    
    if "error" in result:
        print(f"Erreur: {result['error']}")
    else:
        print(f"Nombre total de placeholders trouvés: {result['total_found']}")
        print("Placeholders uniques:")
        for ph in result["unique_placeholders"]:
            print(f"- ${{{ph}}}")
