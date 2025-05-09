import re
from docx import Document
import os

def extract_placeholders_with_context(docx_path):
    """
    Extrait les placeholders avec leur contexte d'un document Word
    
    Returns:
        list: Liste de dictionnaires contenant les placeholders et leur contexte
    """
    try:
        # Vérifier si le fichier existe
        if not os.path.exists(docx_path):
            return {"error": f"Le fichier {docx_path} n'existe pas"}
            
        # Ouvrir le document Word
        doc = Document(docx_path)
        
        # Liste pour stocker tous les placeholders avec contexte
        placeholders_with_context = []
        
        # Garder une trace de la section actuelle
        current_section = "Début du document"
        
        # Pour suivre l'indice du paragraphe
        paragraph_index = 0
        
        # Parcourir tous les paragraphes
        for paragraph in doc.paragraphs:
            paragraph_index += 1
            text = paragraph.text.strip()
            
            # Si c'est un titre, mettre à jour la section actuelle
            if paragraph.style.name.startswith('Heading'):
                current_section = text
                continue
                
            # Chercher les placeholders dans ce paragraphe
            matches = re.finditer(r'\${([^}]*)}', text)
            
            for match in matches:
                placeholder = match.group(1)  # Le contenu entre ${ et }
                start_pos = match.start()
                end_pos = match.end()
                
                # Extraire le contexte (texte avant et après)
                context_before = text[max(0, start_pos-30):start_pos]
                context_after = text[end_pos:min(len(text), end_pos+30)]
                
                # Ajouter à notre liste
                placeholders_with_context.append({
                    "placeholder": placeholder,
                    "full_match": match.group(0),  # ${ et } inclus
                    "paragraph_text": text,
                    "section": current_section,
                    "paragraph_index": paragraph_index,
                    "context_before": context_before,
                    "context_after": context_after,
                    "location": "paragraph",
                    "style": paragraph.style.name
                })
        
        # Parcourir tous les tableaux
        table_index = 0
        for table in doc.tables:
            table_index += 1
            row_index = 0
            for row in table.rows:
                row_index += 1
                cell_index = 0
                for cell in row.cells:
                    cell_index += 1
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        
                        # Chercher les placeholders dans cette cellule
                        matches = re.finditer(r'\${([^}]*)}', text)
                        
                        for match in matches:
                            placeholder = match.group(1)
                            
                            # Ajouter à notre liste avec contexte de tableau
                            placeholders_with_context.append({
                                "placeholder": placeholder,
                                "full_match": match.group(0),
                                "paragraph_text": text,
                                "section": current_section,
                                "location": "table",
                                "table_index": table_index,
                                "row_index": row_index,
                                "cell_index": cell_index
                            })
        
        # Créer un dictionnaire des placeholders uniques avec leurs occurrences
        unique_placeholders = {}
        for item in placeholders_with_context:
            ph = item["placeholder"]
            if ph not in unique_placeholders:
                unique_placeholders[ph] = {
                    "placeholder": ph,
                    "occurrences": 1,
                    "contexts": [item]
                }
            else:
                unique_placeholders[ph]["occurrences"] += 1
                unique_placeholders[ph]["contexts"].append(item)
        
        return {
            "total_found": len(placeholders_with_context),
            "unique_count": len(unique_placeholders),
            "unique_placeholders": list(unique_placeholders.values())
        }
        
    except Exception as e:
        return {"error": str(e)}

def organize_placeholders_by_section(placeholders_data):
    """
    Organise les placeholders par section du document
    """
    sections = {}
    
    for ph in placeholders_data["unique_placeholders"]:
        for context in ph["contexts"]:
            section = context["section"]
            
            if section not in sections:
                sections[section] = []
                
            sections[section].append({
                "placeholder": ph["placeholder"],
                "context": context["paragraph_text"][:50] + "..." if len(context["paragraph_text"]) > 50 else context["paragraph_text"]
            })
    
    return {
        "total_sections": len(sections),
        "sections": [{"name": k, "placeholders": v} for k, v in sections.items()]
    }

def analyze_placeholder_types(placeholders_data):
    """
    Essaie de déterminer le type de données attendu pour chaque placeholder
    """
    type_indicators = {
        "date": ["date", "jour", "mois", "année", "calendar"],
        "number": ["nombre", "montant", "quantité", "pourcentage", "total", "somme"],
        "person": ["nom", "prénom", "personne", "contact", "responsable"],
        "boolean": ["oui/non", "vrai/faux", "choix"],
        "address": ["adresse", "ville", "code postal", "pays"],
        "email": ["email", "courriel", "mail"],
        "phone": ["téléphone", "tel", "mobile", "fixe"],
    }
    
    for ph in placeholders_data["unique_placeholders"]:
        placeholder = ph["placeholder"].lower()
        
        # Déterminer le type en fonction du nom du placeholder
        detected_type = "text"  # type par défaut
        
        for type_name, indicators in type_indicators.items():
            if any(indicator in placeholder for indicator in indicators):
                detected_type = type_name
                break
        
        # Ajouter le type détecté au placeholder
        ph["detected_type"] = detected_type
        
    return placeholders_data

# Test de la fonction
if __name__ == "__main__":
    # Remplacer par le chemin de votre template
    template_path = "chemin/vers/votre/template.docx"
    result = extract_placeholders_with_context(template_path)
    
    if "error" in result:
        print(f"Erreur: {result['error']}")
    else:
        print(f"Nombre total de placeholders trouvés: {result['total_found']}")
        print(f"Nombre de placeholders uniques: {result['unique_count']}")
        
        # Afficher les 3 premiers placeholders avec leur contexte
        for i, ph in enumerate(result["unique_placeholders"][:3]):
            print(f"\nPlaceholder {i+1}: ${{{ph['placeholder']}}}")
            print(f"Nombre d'occurrences: {ph['occurrences']}")
            print("Premier contexte:")
            ctx = ph['contexts'][0]
            print(f"  Section: {ctx['section']}")
            print(f"  Texte: ...{ctx['context_before']}${{{ph['placeholder']}}}{ctx['context_after']}...")
